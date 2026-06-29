import sys
import logging
import subprocess
from pathlib import Path

log = logging.getLogger("ExtractorAgent")

try:
    import pdfplumber
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pdfplumber"])
    import pdfplumber

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx


def extract_clinical_data(file_path: str) -> dict:
    path = Path(file_path)
    log.info(f"[Extractor] Starting extraction — file: {path.name}")

    if not path.exists():
        log.error(f"[Extractor] File not found: {file_path}")
        return {"text": "", "tables": [], "status": "failed", "error": "File not found"}

    ext = path.suffix.lower()
    log.info(f"[Extractor] File type: {ext} | Size: {path.stat().st_size} bytes")

    extracted_text   = ""
    extracted_tables = []

    try:
        if ext == ".pdf":
            log.info("[Extractor] Parsing PDF with pdfplumber…")
            with pdfplumber.open(path) as pdf:
                log.info(f"[Extractor] PDF has {len(pdf.pages)} page(s)")
                for page_num, page in enumerate(pdf.pages, start=1):
                    text_content = page.extract_text()
                    if text_content:
                        extracted_text += f"\n--- Page {page_num} ---\n{text_content}"
                        log.debug(f"[Extractor] Page {page_num}: {len(text_content)} chars")

                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            extracted_tables.append({"source_page": page_num, "data": table})
                            log.debug(f"[Extractor] Page {page_num}: found table with {len(table)} rows")

        elif ext == ".docx":
            log.info("[Extractor] Parsing DOCX…")
            doc = docx.Document(path)
            extracted_text = "\n".join([p.text for p in doc.paragraphs if p.text])
            for i, table in enumerate(doc.tables):
                table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if table_data:
                    extracted_tables.append({"source_table_index": i, "data": table_data})
            log.info(f"[Extractor] DOCX: {len(extracted_text)} chars, {len(extracted_tables)} tables")

        elif ext in (".txt", ".log"):
            log.info("[Extractor] Reading plain text file…")
            with open(path, "r", encoding="utf-8") as f:
                extracted_text = f.read()
            log.info(f"[Extractor] Text: {len(extracted_text)} chars")

        elif ext == ".json":
            import json
            log.info("[Extractor] Reading JSON file…")
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                extracted_text = json.dumps(data, indent=2)
            log.info(f"[Extractor] JSON: {len(extracted_text)} chars")

        elif ext in (".png", ".jpg", ".jpeg"):
            log.info("[Extractor] Running OCR via pytesseract…")
            try:
                import pytesseract
                from PIL import Image
            except ImportError:
                subprocess.check_call([sys.executable, "-m", "pip", "install", "pytesseract", "Pillow"])
                import pytesseract
                from PIL import Image
            extracted_text = pytesseract.image_to_string(Image.open(path))
            log.info(f"[Extractor] OCR result: {len(extracted_text)} chars")

        else:
            log.warning(f"[Extractor] Unsupported file type: {ext}")
            return {"text": "", "tables": [], "status": "unsupported_format"}

        log.info(
            f"[Extractor] Done — {len(extracted_text)} chars, "
            f"{len(extracted_tables)} tables extracted from {path.name}"
        )
        return {
            "text"     : extracted_text.strip(),
            "tables"   : extracted_tables,
            "status"   : "success",
            "file_name": path.name,
        }

    except Exception as e:
        log.exception(f"[Extractor] FAILED parsing {path.name}: {e}")
        return {"text": "", "tables": [], "status": "error", "error": str(e)}
